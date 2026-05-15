"""
Document parsing utilities for Contract Intelligence Agent
Supports PDF, DOCX, and TXT files
"""
import io
import logging
from pathlib import Path
from typing import Optional, Tuple
import PyPDF2
import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)


class DocumentParsingError(Exception):
    """Custom exception for document parsing errors"""
    pass


class DocumentParser:
    """
    Unified document parser for multiple file formats
    """
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt'}
    
    @staticmethod
    def parse_file(file_content: bytes, filename: str) -> Tuple[str, str]:
        """
        Parse document and extract text
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            
        Returns:
            Tuple of (extracted_text, file_type)
            
        Raises:
            DocumentParsingError: If parsing fails
        """
        file_extension = Path(filename).suffix.lower()
        
        if file_extension not in DocumentParser.SUPPORTED_EXTENSIONS:
            raise DocumentParsingError(
                f"Unsupported file type: {file_extension}. "
                f"Supported types: {', '.join(DocumentParser.SUPPORTED_EXTENSIONS)}"
            )
        
        try:
            if file_extension == '.pdf':
                text = DocumentParser._parse_pdf(file_content)
                file_type = 'pdf'
            elif file_extension in {'.docx', '.doc'}:
                text = DocumentParser._parse_docx(file_content)
                file_type = 'docx'
            elif file_extension == '.txt':
                text = DocumentParser._parse_txt(file_content)
                file_type = 'txt'
            else:
                raise DocumentParsingError(f"Unsupported file extension: {file_extension}")
            
            # Validate extracted text
            if not text or len(text.strip()) < 100:
                raise DocumentParsingError(
                    "Extracted text is too short or empty. "
                    "Please ensure the document contains readable text."
                )
            
            # Clean and normalize text
            text = DocumentParser._clean_text(text)
            
            logger.info(f"Successfully parsed {filename} ({file_type}), extracted {len(text)} characters")
            return text, file_type
            
        except Exception as e:
            logger.error(f"Error parsing {filename}: {str(e)}")
            raise DocumentParsingError(f"Failed to parse document: {str(e)}")
    
    @staticmethod
    def _parse_pdf(file_content: bytes) -> str:
        """
        Parse PDF file using multiple methods for best results
        
        Args:
            file_content: PDF file bytes
            
        Returns:
            Extracted text
        """
        text_parts = []
        
        # Method 1: Try pdfplumber (better for complex layouts)
        try:
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            if text_parts:
                logger.info("Successfully extracted PDF using pdfplumber")
                return "\n\n".join(text_parts)
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {e}, trying PyPDF2")
        
        # Method 2: Fallback to PyPDF2
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            if text_parts:
                logger.info("Successfully extracted PDF using PyPDF2")
                return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"PyPDF2 extraction failed: {e}")
            raise DocumentParsingError("Failed to extract text from PDF")
        
        raise DocumentParsingError("No text could be extracted from PDF")
    
    @staticmethod
    def _parse_docx(file_content: bytes) -> str:
        """
        Parse DOCX file
        
        Args:
            file_content: DOCX file bytes
            
        Returns:
            Extracted text
        """
        try:
            doc = Document(io.BytesIO(file_content))
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            if not text_parts:
                raise DocumentParsingError("No text found in DOCX file")
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"DOCX parsing error: {e}")
            raise DocumentParsingError(f"Failed to parse DOCX: {str(e)}")
    
    @staticmethod
    def _parse_txt(file_content: bytes) -> str:
        """
        Parse plain text file
        
        Args:
            file_content: Text file bytes
            
        Returns:
            Extracted text
        """
        try:
            # Try UTF-8 first
            try:
                text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                # Fallback to latin-1
                text = file_content.decode('latin-1')
            
            if not text.strip():
                raise DocumentParsingError("Text file is empty")
            
            return text
            
        except Exception as e:
            logger.error(f"TXT parsing error: {e}")
            raise DocumentParsingError(f"Failed to parse TXT: {str(e)}")
    
    @staticmethod
    def _clean_text(text: str) -> str:
        """
        Clean and normalize extracted text
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        lines = [line.strip() for line in text.split('\n')]
        lines = [line for line in lines if line]
        
        # Join with proper spacing
        cleaned = '\n'.join(lines)
        
        # Remove multiple consecutive blank lines
        while '\n\n\n' in cleaned:
            cleaned = cleaned.replace('\n\n\n', '\n\n')
        
        return cleaned.strip()
    
    @staticmethod
    def validate_file_size(file_size: int, max_size: int = 10 * 1024 * 1024) -> None:
        """
        Validate file size
        
        Args:
            file_size: File size in bytes
            max_size: Maximum allowed size in bytes (default 10MB)
            
        Raises:
            DocumentParsingError: If file is too large
        """
        if file_size > max_size:
            raise DocumentParsingError(
                f"File size ({file_size / 1024 / 1024:.2f}MB) exceeds "
                f"maximum allowed size ({max_size / 1024 / 1024:.2f}MB)"
            )
    
    @staticmethod
    def validate_file_extension(filename: str) -> None:
        """
        Validate file extension
        
        Args:
            filename: File name
            
        Raises:
            DocumentParsingError: If extension is not supported
        """
        extension = Path(filename).suffix.lower()
        if extension not in DocumentParser.SUPPORTED_EXTENSIONS:
            raise DocumentParsingError(
                f"Unsupported file type: {extension}. "
                f"Supported types: {', '.join(DocumentParser.SUPPORTED_EXTENSIONS)}"
            )


# Made with Bob